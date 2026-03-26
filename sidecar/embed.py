#!/usr/bin/env python3
"""Embedding sidecar for tendrl-engine.

Runs a small HTTP server that generates text embeddings using
sentence-transformers. The Rust engine calls this to embed event
content and search queries.

Usage:
    python embed.py [--model MODEL] [--port PORT]
    MODEL=all-MiniLM-L6-v2 PORT=3031 python embed.py
"""

import argparse
import os
import sys

from flask import Flask, jsonify, request
from sentence_transformers import SentenceTransformer

app = Flask(__name__)

model = None
model_name = None


def get_model():
    global model, model_name
    if model is None:
        name = os.environ.get("MODEL", "all-MiniLM-L6-v2")
        print(f"Loading model: {name}", file=sys.stderr)
        model = SentenceTransformer(name)
        model_name = name
        print(
            f"Model loaded: {name} ({model.get_sentence_embedding_dimension()}d)",
            file=sys.stderr,
        )
    return model


@app.route("/health", methods=["GET"])
def health():
    m = get_model()
    return jsonify(
        {
            "status": "ok",
            "model": model_name,
            "dimensions": m.get_sentence_embedding_dimension(),
        }
    )


@app.route("/embed", methods=["POST"])
def embed():
    data = request.get_json(force=True)
    texts = data.get("texts", [])
    if not texts:
        return jsonify({"vectors": []})

    m = get_model()

    # Sub-batch to avoid OOM on large requests
    batch_size = 512
    all_vectors = []
    for i in range(0, len(texts), batch_size):
        batch = texts[i : i + batch_size]
        embeddings = m.encode(batch, normalize_embeddings=True)
        all_vectors.extend(embeddings.tolist())

    return jsonify({"vectors": all_vectors})


# ── Document parsing ──────────────────────────────────────────────


def parse_pdf(file_bytes, filename):
    import fitz  # pymupdf

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if text:
            pages.append({"page_num": i + 1, "title": None, "content": text})
    doc.close()
    return {"filename": filename, "format": "pdf", "page_count": len(pages), "pages": pages}


def parse_docx(file_bytes, filename):
    import io
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    pages = []
    current_title = None
    current_content = []
    page_num = 0

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            # Save previous section
            if current_content:
                page_num += 1
                pages.append({
                    "page_num": page_num,
                    "title": current_title,
                    "content": "\n".join(current_content).strip(),
                })
                current_content = []
            current_title = para.text.strip() or None
        else:
            text = para.text.strip()
            if text:
                current_content.append(text)

    # Last section
    if current_content:
        page_num += 1
        pages.append({
            "page_num": page_num,
            "title": current_title,
            "content": "\n".join(current_content).strip(),
        })

    # Fallback: if no headings found, treat whole doc as one page
    if not pages:
        full_text = "\n".join(p.text for p in doc.paragraphs).strip()
        if full_text:
            pages = [{"page_num": 1, "title": None, "content": full_text}]

    return {"filename": filename, "format": "docx", "page_count": len(pages), "pages": pages}


def parse_text(file_bytes, filename):
    """Parse plain text, markdown, org-mode, or asciidoc by headings."""
    text = file_bytes.decode("utf-8", errors="replace")
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

    # Detect heading pattern
    import re
    if ext == "org":
        pattern = re.compile(r"^(\*+)\s+(.+)$", re.MULTILINE)
    elif ext in ("adoc", "asciidoc"):
        pattern = re.compile(r"^(=+)\s+(.+)$", re.MULTILINE)
    else:  # md, txt
        pattern = re.compile(r"^(#+)\s+(.+)$", re.MULTILINE)

    splits = list(pattern.finditer(text))

    if not splits:
        # No headings — return whole file as one page
        return {
            "filename": filename,
            "format": ext,
            "page_count": 1,
            "pages": [{"page_num": 1, "title": None, "content": text.strip()}],
        }

    pages = []
    for i, match in enumerate(splits):
        title = match.group(2).strip()
        start = match.end()
        end = splits[i + 1].start() if i + 1 < len(splits) else len(text)
        content = text[start:end].strip()
        if content or title:
            pages.append({"page_num": i + 1, "title": title, "content": content})

    return {"filename": filename, "format": ext, "page_count": len(pages), "pages": pages}


def parse_epub(file_bytes, filename):
    import io
    import ebooklib
    from ebooklib import epub
    from lxml import etree

    book = epub.read_epub(io.BytesIO(file_bytes))
    pages = []
    page_num = 0

    for item in book.get_items_of_type(ebooklib.ITEM_DOCUMENT):
        try:
            tree = etree.fromstring(item.get_content())
            text = etree.tostring(tree, method="text", encoding="unicode").strip()
            if text and len(text) > 20:
                page_num += 1
                title = item.get_name().rsplit("/", 1)[-1].rsplit(".", 1)[0]
                pages.append({"page_num": page_num, "title": title, "content": text})
        except Exception:
            continue

    return {"filename": filename, "format": "epub", "page_count": len(pages), "pages": pages}


def parse_html(file_bytes, filename):
    from lxml import etree
    import re

    text = file_bytes.decode("utf-8", errors="replace")
    try:
        tree = etree.HTML(text)
        body = tree.find(".//body")
        if body is None:
            body = tree
        content = etree.tostring(body, method="text", encoding="unicode").strip()
        content = re.sub(r"\n{3,}", "\n\n", content)
    except Exception:
        content = text

    return {
        "filename": filename,
        "format": "html",
        "page_count": 1,
        "pages": [{"page_num": 1, "title": None, "content": content}],
    }


PARSERS = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "epub": parse_epub,
    "html": parse_html,
    "htm": parse_html,
    "txt": parse_text,
    "md": parse_text,
    "org": parse_text,
    "adoc": parse_text,
    "asciidoc": parse_text,
    "rst": parse_text,
}


@app.route("/parse", methods=["POST"])
def parse_document():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    f = request.files["file"]
    filename = f.filename or "unknown"
    file_bytes = f.read()
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    parser = PARSERS.get(ext)
    if parser is None:
        return jsonify({"error": f"Unsupported format: .{ext}"}), 400

    try:
        result = parser(file_bytes, filename)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Embedding sidecar")
    parser.add_argument(
        "--model", default=os.environ.get("MODEL", "all-MiniLM-L6-v2")
    )
    parser.add_argument(
        "--port", type=int, default=int(os.environ.get("PORT", "3031"))
    )
    parser.add_argument("--host", default="127.0.0.1")
    args = parser.parse_args()

    os.environ["MODEL"] = args.model

    # Preload model
    get_model()

    print(f"Embedding sidecar listening on http://{args.host}:{args.port}")
    app.run(host=args.host, port=args.port, debug=False)
